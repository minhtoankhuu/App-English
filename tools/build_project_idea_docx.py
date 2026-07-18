from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/product/ENGLISH_EXAM_AI_PRODUCT_REQUIREMENTS.vi.md"
OUTPUT = ROOT / "docs/product/ENGLISH_EXAM_AI_PRODUCT_REQUIREMENTS.vi.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "172B4D"
LIGHT = "F2F4F7"
BORDER = "CBD5E1"
MUTED = "667085"


def set_run_font(run, name="Aptos", size=11, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def mark_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    settings = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in settings.items():
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    if "Metadata" not in doc.styles:
        style = doc.styles.add_style("Metadata", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.space_after = Pt(3)


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("ĐẶC TẢ Ý TƯỞNG SẢN PHẨM • 2026")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(78)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("ĐẶC TẢ Ý TƯỞNG\nSẢN PHẨM")
    set_run_font(r, "Aptos Display", 28, True, BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("Nền tảng AI tạo bài tập và đề tiếng Anh\ntheo kiến thức sách giáo khoa")
    set_run_font(r, "Aptos Display", 19, True, INK)

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Cm(3.2), Cm(12.8)]
    values = [
        ("Phiên bản", "1.3"),
        ("Ngày", "18/07/2026"),
        ("Trạng thái", "Thiết kế đã thống nhất trước khi lập kế hoạch triển khai"),
    ]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
            set_cell_margins(cell, 100, 140, 100, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text = values[i][0]
        row.cells[1].text = values[i][1]
        set_cell_shading(row.cells[0], LIGHT)
        for run in row.cells[0].paragraphs[0].runs:
            set_run_font(run, size=10, bold=True, color=DARK_BLUE)
        for run in row.cells[1].paragraphs[0].runs:
            set_run_font(run, size=10, color=INK)
    set_table_borders(table)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Định hướng cốt lõi")
    set_run_font(r, size=10, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.right_indent = Cm(0.45)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Giáo viên kiểm soát cấu trúc và chất lượng; AI hỗ trợ sinh nội dung có nguồn; DOCX là đầu ra có thể tiếp tục chỉnh sửa và in.")
    set_run_font(r, size=11, bold=True, color=INK)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EAF2F8")
    p_pr.append(shd)

    doc.add_page_break()


def add_inline_text(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 9.5, False, DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_content(doc, text):
    lines = text.splitlines()
    i = 0
    skip_front = True
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if skip_front:
            if line == "## 1. Tóm tắt điều hành":
                skip_front = False
            else:
                i += 1
                continue
        if not line:
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            total = Cm(16.6)
            col_widths = [Cm(4.3), Cm(12.3)] if len(table.columns) == 2 else [total / len(table.columns)] * len(table.columns)
            for ri, row in enumerate(rows):
                prevent_row_split(table.rows[ri])
                if ri == 0:
                    mark_repeat_table_header(table.rows[ri])
                for ci, value in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.width = col_widths[ci]
                    cell.text = value
                    set_cell_margins(cell)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if ri == 0:
                        set_cell_shading(cell, LIGHT)
                    for run in cell.paragraphs[0].runs:
                        set_run_font(run, size=9.5, bold=(ri == 0), color=DARK_BLUE if ri == 0 else INK)
            set_table_borders(table)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("#### "):
            doc.add_paragraph(line[5:], style="Heading 3")
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_text(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_text(p, line[2:])
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line.replace("**", ""))
            set_run_font(r, "Aptos Display", 11, True, BLUE)
        else:
            p = doc.add_paragraph()
            add_inline_text(p, line)
        i += 1


def main():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    core = doc.core_properties
    core.title = "Đặc tả ý tưởng sản phẩm AI tạo đề tiếng Anh"
    core.subject = "Product requirements và thiết kế khái niệm"
    core.author = "Codex và chủ dự án"
    core.keywords = "AI, RAG, English exam generator, DOCX, template"
    add_cover(doc)
    add_markdown_content(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
