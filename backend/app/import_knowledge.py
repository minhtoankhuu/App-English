"""Import tài liệu Global Success (G6-G8) vào knowledge_documents/knowledge_chunks.

Idempotent theo checksum SHA-256: file không đổi thì bỏ qua, file đổi nội dung
thì xoá toàn bộ chunk cũ và parse lại. Không đụng Cambridge/Tense/G9 (khác cấu
trúc — xem docs/superpowers/specs/2026-07-19-knowledge-base-global-success-design.md).

Chạy độc lập: `python -m app.import_knowledge` — không gộp vào seed.py vì đây là
nhập tài liệu lớn, không phải danh mục tĩnh (tránh làm chậm mọi lần khởi động).
"""

import argparse
import hashlib
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.config import get_settings
from app.db import SessionLocal
from app.models.academic import Grade, Unit
from app.models.knowledge import DocumentChunkType, KnowledgeChunk, KnowledgeDocument
from app.services.docx_utils import paragraph_markup
from app.services.exam_parser import parse_exam_items
from app.services.knowledge_parser import parse_lesson_docx
from app.services.knowledge_parser_g9 import parse_g9_vocabulary

GRADES = (6, 7, 8)
_UNIT_NUMBER_RE = re.compile(r"UNIT\s*(\d+)", re.IGNORECASE)
# G9 khác cấu trúc: 1 file gộp 12 Unit (xem knowledge_parser_g9.py).
_G9_VOCAB_FILE = "Vocabulary Global Success 9 (1).docx"


@dataclass
class ImportStats:
    files_seen: int = 0
    documents_created: int = 0
    documents_updated: int = 0
    documents_unchanged: int = 0
    chunks_written: int = 0


def _checksum(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _unit_number(file_name: str) -> int | None:
    match = _UNIT_NUMBER_RE.search(file_name)
    return int(match.group(1)) if match else None


def import_global_success(db: Session, base_path: Path, force: bool = False) -> ImportStats:
    """`base_path` là gốc Knowledge_Base/ — file thật nằm dưới `Global Success/G{6,7,8}/`.

    `force=True` bỏ qua so khớp checksum, parse lại toàn bộ file đã import — dùng khi
    sửa logic parser (`docx_utils`/`knowledge_parser`) mà nội dung file .docx không đổi,
    nên checksum vẫn khớp và bản ghi cũ (chunk cũ) sẽ không tự refresh nếu không có cờ này."""
    stats = ImportStats()
    grades = {g.number: g for g in db.scalars(select(Grade).where(Grade.number.in_(GRADES)))}
    global_success_dir = base_path / "Global Success"

    for grade_number in GRADES:
        grade = grades.get(grade_number)
        folder = global_success_dir / f"G{grade_number}"
        if grade is None or not folder.is_dir():
            continue

        units_by_order = {u.order_no: u for u in db.scalars(select(Unit).where(Unit.grade_id == grade.id))}

        for file_path in sorted(folder.glob("*.docx")):
            stats.files_seen += 1
            order_no = _unit_number(file_path.name)
            unit = units_by_order.get(order_no) if order_no else None
            if unit is None:
                continue

            checksum = _checksum(file_path)
            existing = db.scalar(
                select(KnowledgeDocument).where(
                    KnowledgeDocument.unit_id == unit.id,
                    KnowledgeDocument.file_name == file_path.name,
                )
            )

            if existing is not None and existing.checksum == checksum and not force:
                stats.documents_unchanged += 1
                continue

            parsed_chunks = parse_lesson_docx(file_path)

            if existing is not None:
                existing.checksum = checksum
                existing.chunks.clear()
                document = existing
                stats.documents_updated += 1
            else:
                document = KnowledgeDocument(unit_id=unit.id, file_name=file_path.name, checksum=checksum)
                db.add(document)
                stats.documents_created += 1

            for chunk in parsed_chunks:
                document.chunks.append(
                    KnowledgeChunk(
                        order_no=chunk.order_no,
                        chunk_type=chunk.chunk_type,
                        section_title=chunk.section_title,
                        raw_text=chunk.raw_text,
                        structured=chunk.structured,
                    )
                )
            stats.chunks_written += len(parsed_chunks)
            db.flush()

    return stats


def import_exam_papers(db: Session, base_path: Path, force: bool = False) -> ImportStats:
    """Nhập đề thi thật từ `Knowledge_Base/Exams/G{6..9}/*.docx` thành chunk EXAM_ITEM.

    Khác với sách giáo khoa (mỗi chunk là một mục từ), mỗi chunk ở đây là MỘT CÂU HỎI
    hoàn chỉnh giữ nguyên như đề in ra — vì thứ model cần bắt chước là cả câu. Sách
    Global Success gần như không có câu ví dụ (GS7 Unit 1: 2/231 đoạn) nên đây là nguồn
    câu mẫu thật duy nhất. Idempotent theo checksum như import_global_success.
    """
    stats = ImportStats()
    exams_dir = base_path / "Exams"
    if not exams_dir.is_dir():
        return stats

    grades = {g.number: g for g in db.scalars(select(Grade))}

    for grade_number, grade in sorted(grades.items()):
        folder = exams_dir / f"G{grade_number}"
        if not folder.is_dir():
            continue
        units_by_order = {u.order_no: u for u in db.scalars(select(Unit).where(Unit.grade_id == grade.id))}

        for file_path in sorted(folder.glob("*.docx")):
            stats.files_seen += 1
            order_no = _unit_number(file_path.name)
            unit = units_by_order.get(order_no) if order_no else None
            if unit is None:
                continue

            checksum = _checksum(file_path)
            existing = db.scalar(
                select(KnowledgeDocument).where(
                    KnowledgeDocument.unit_id == unit.id,
                    KnowledgeDocument.file_name == file_path.name,
                )
            )
            if existing is not None and existing.checksum == checksum and not force:
                stats.documents_unchanged += 1
                continue

            # paragraph_markup giữ gạch chân của đề thật — nội dung chính của câu phát âm.
            items = parse_exam_items(
                [paragraph_markup(p) for p in Document(str(file_path)).paragraphs]
            )

            if existing is not None:
                existing.checksum = checksum
                existing.chunks.clear()
                document = existing
                stats.documents_updated += 1
            else:
                document = KnowledgeDocument(unit_id=unit.id, file_name=file_path.name, checksum=checksum)
                db.add(document)
                stats.documents_created += 1

            for order, item in enumerate(items, start=1):
                document.chunks.append(
                    KnowledgeChunk(
                        order_no=order,
                        chunk_type=DocumentChunkType.EXAM_ITEM,
                        # Mã dạng bài nằm ở section_title để truy xuất lọc được câu mẫu
                        # đúng dạng đang sinh (xem rag_search.exam_examples).
                        section_title=item.exercise_type_code,
                        raw_text=item.text,
                        structured=None,
                    )
                )
            stats.chunks_written += len(items)
            db.flush()

    return stats


def import_grade9_vocabulary(db: Session, base_path: Path, force: bool = False) -> ImportStats:
    """Nhập từ vựng Global Success 9 từ file gộp 12 Unit. Mỗi Unit -> 1 KnowledgeDocument
    (unit_id theo order_no), file_name chung nên checksum giống nhau cho mọi Unit —
    idempotent theo checksum như import_global_success."""
    stats = ImportStats()
    grade9 = db.scalar(select(Grade).where(Grade.number == 9))
    vocab_file = base_path / "Global Success" / "G9" / _G9_VOCAB_FILE
    if grade9 is None or not vocab_file.is_file():
        return stats

    units_by_order = {u.order_no: u for u in db.scalars(select(Unit).where(Unit.grade_id == grade9.id))}
    checksum = _checksum(vocab_file)
    per_unit = parse_g9_vocabulary(vocab_file)

    for order_no, parsed_chunks in per_unit.items():
        unit = units_by_order.get(order_no)
        if unit is None:
            continue
        stats.files_seen += 1
        existing = db.scalar(
            select(KnowledgeDocument).where(
                KnowledgeDocument.unit_id == unit.id,
                KnowledgeDocument.file_name == _G9_VOCAB_FILE,
            )
        )
        if existing is not None and existing.checksum == checksum and not force:
            stats.documents_unchanged += 1
            continue

        if existing is not None:
            existing.checksum = checksum
            existing.chunks.clear()
            document = existing
            stats.documents_updated += 1
        else:
            document = KnowledgeDocument(unit_id=unit.id, file_name=_G9_VOCAB_FILE, checksum=checksum)
            db.add(document)
            stats.documents_created += 1

        for chunk in parsed_chunks:
            document.chunks.append(
                KnowledgeChunk(
                    order_no=chunk.order_no,
                    chunk_type=chunk.chunk_type,
                    section_title=chunk.section_title,
                    raw_text=chunk.raw_text,
                    structured=chunk.structured,
                )
            )
        stats.chunks_written += len(parsed_chunks)
        db.flush()

    return stats


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--force",
        action="store_true",
        help="Parse lại toàn bộ file kể cả khi checksum không đổi (dùng khi logic parser vừa sửa).",
    )
    args = parser.parse_args()

    settings = get_settings()
    db = SessionLocal()
    try:
        base_path = Path(settings.knowledge_base_dir)
        stats = import_global_success(db, base_path, force=args.force)
        g9_stats = import_grade9_vocabulary(db, base_path, force=args.force)
        exam_stats = import_exam_papers(db, base_path, force=args.force)
        db.commit()
        print(
            f"Import OK (G6-8): {stats.files_seen} file, {stats.documents_created} mới, "
            f"{stats.documents_updated} cập nhật, {stats.documents_unchanged} không đổi, "
            f"{stats.chunks_written} chunk ghi."
        )
        print(
            f"Import OK (G9 vocab): {g9_stats.files_seen} unit, {g9_stats.documents_created} mới, "
            f"{g9_stats.documents_updated} cập nhật, {g9_stats.documents_unchanged} không đổi, "
            f"{g9_stats.chunks_written} chunk ghi."
        )
        print(
            f"Import OK (de thi that): {exam_stats.files_seen} file, {exam_stats.documents_created} mới, "
            f"{exam_stats.documents_updated} cập nhật, {exam_stats.documents_unchanged} không đổi, "
            f"{exam_stats.chunks_written} câu mẫu ghi."
        )
    finally:
        db.close()


if __name__ == "__main__":
    main()
