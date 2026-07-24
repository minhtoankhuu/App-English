"""Dựng file DOCX từ đề đã duyệt, theo đúng thông số đã chốt với giáo viên
(Implementation Notes mục 2 / PRD mục 14). Không dùng trực tiếp văn bản thô của
LLM — dữ liệu đến từ Question đã qua Validation Engine và được giáo viên duyệt."""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi.responses import StreamingResponse

from app.models.exam import Exam, ExamVariant, ExportMode
from app.services.text_markup import UNDERLINE_MARKUP_RE, dedupe_pronunciation_suffix

RED = RGBColor(0xC0, 0x00, 0x00)
FONT_NAME = "Times New Roman"
# LLM đánh dấu phần cần gạch chân bằng <u>...</u> trong text lựa chọn (dạng phát âm/
# trọng âm — PRD 7.3 yêu cầu gạch chân âm/trọng âm khác biệt trong đề in ra, xem
# app/services/prompts.py). Không dùng docx run.underline trực tiếp từ LLM vì output
# LLM là JSON text thuần, không phải OOXML.
_UNDERLINE_MARKUP_RE = UNDERLINE_MARKUP_RE
PAGE_WIDTH_CM = 21.0
MARGIN_CM = 1.27
USABLE_WIDTH_CM = PAGE_WIDTH_CM - 2 * MARGIN_CM
ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
# Ngưỡng để chuyển 4 lựa chọn/dòng xuống 2 lựa chọn/dòng. Cột 4 lựa chọn rộng
# ~4.6cm — để dư biên an toàn (chữ đậm/gạch chân của đáp án tô đỏ rộng hơn chữ
# thường) vì lựa chọn sát ngưỡng cũ (24) từng bị tràn cột, kéo cả dòng xuống
# dòng dưới trong Word (báo cáo giáo viên 21/07/2026).
LONG_OPTION_THRESHOLD = 18
# Thụt lề tiêu đề + nội dung câu hỏi thuộc 1 Phần con (vd "1. Đuôi -s/-es" trong
# khối "I. PRONUNCIATION") để phân cấp rõ I. > 1. > câu hỏi, khớp mẫu đề tham khảo
# (chủ dự án yêu cầu thụt sâu hơn mức 0.5cm ban đầu — 21/07/2026).
PART_CONTENT_INDENT_CM = 0.9
# Khoảng chừa cho số thứ tự câu hỏi ("6.", "11.") khi in chung dòng với lựa chọn
# (xem _render_options) — đủ rộng cho số 2 chữ số + dấu chấm ở cỡ chữ 11.5pt.
NUMBER_GUTTER_CM = 0.9


def _set_font(run, size: float = 11.5, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)


def _add_runs_with_underline(p, text: str, *, bold: bool = False, color: RGBColor | None = None) -> None:
    """Tách `text` theo marker `<u>...</u>` thành nhiều run — phần trong marker được
    `run.font.underline = True`, phần còn lại giữ nguyên. Không có marker thì y hệt
    hành vi cũ (1 run duy nhất)."""
    pos = 0
    for match in _UNDERLINE_MARKUP_RE.finditer(text):
        if match.start() > pos:
            run = p.add_run(text[pos : match.start()])
            _set_font(run, bold=bold, color=color)
        underline_run = p.add_run(match.group(1))
        # Phần gạch chân luôn in đậm để dễ nhận ra (yêu cầu chủ dự án 21/07/2026),
        # bất kể phần còn lại của lựa chọn có đang bold vì là đáp án đúng hay không.
        _set_font(underline_run, bold=True, color=color)
        underline_run.font.underline = True
        pos = match.end()
    if pos < len(text):
        run = p.add_run(text[pos:])
        _set_font(run, bold=bold, color=color)


def _shared_instruction_question_ids(ordered_questions: list) -> set:
    """ID các câu thuộc nhóm dùng CHUNG y hệt 1 câu dẫn (vd dạng phát âm/trọng âm:
    "Which word has a different stress pattern?" lặp cho mọi câu trong nhóm).

    Nhóm như vậy in câu dẫn thành 1 dòng "hướng dẫn" riêng, và MỌI câu đều giữ số thứ
    tự ở dòng lựa chọn — thay vì gộp số vào dòng câu dẫn khiến câu đầu nhóm mất số
    (báo cáo giáo viên 21/07/2026). Câu có câu dẫn riêng (trắc nghiệm thường) không
    thuộc nhóm nào nên giữ nguyên kiểu "số + câu dẫn cùng dòng".
    """
    shared: set = set()
    group: list = []

    def flush(items: list) -> None:
        first = items[0].prompt_text if items else None
        if len(items) > 1 and first and all(q.prompt_text == first for q in items):
            shared.update(q.id for q in items)

    group_part_id = None
    for question in ordered_questions:
        if not group or question.part_id != group_part_id:
            flush(group)
            group = []
            group_part_id = question.part_id
        group.append(question)
    flush(group)
    return shared


def _new_paragraph(doc: Document, *, justify: bool = False, center: bool = False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _visible_length(text: str) -> int:
    """Độ dài không tính marker <u></u> (14 ký tự không hiển thị) — tránh option ngắn
    có gạch chân bị tính nhầm là dài, kéo cả hàng xuống 2 lựa chọn/dòng."""
    return len(_UNDERLINE_MARKUP_RE.sub(r"\1", text))


def _render_options(
    doc: Document, options: list[dict], with_key: bool, *, indent_cm: float = 0.0, number: str | None = None
) -> None:
    """`number` (vd "6.") in ngay đầu dòng lựa chọn đầu tiên, cùng 1 dòng — dùng khi
    câu không có câu dẫn riêng để hiện (đã in ở câu trước đó cùng phần/khối), tránh
    số thứ tự đứng 1 mình 1 dòng rồi lựa chọn "xuống dòng" tách rời bên dưới (báo
    cáo giáo viên 21/07/2026)."""
    max_len = max(_visible_length(opt["text"]) for opt in options)
    per_line = 2 if max_len > LONG_OPTION_THRESHOLD else 4
    gutter_cm = NUMBER_GUTTER_CM if number else 0.0
    step_cm = (USABLE_WIDTH_CM - indent_cm - gutter_cm) / per_line

    for i in range(0, len(options), per_line):
        row = options[i : i + per_line]
        p = _new_paragraph(doc)
        if indent_cm:
            p.paragraph_format.left_indent = Cm(indent_cm)
        if gutter_cm:
            p.paragraph_format.tab_stops.add_tab_stop(Cm(gutter_cm))
            if i == 0:
                no_run = p.add_run(number)
                _set_font(no_run, bold=True)
            p.add_run().add_tab()
        for k in range(1, len(row)):
            p.paragraph_format.tab_stops.add_tab_stop(Cm(gutter_cm + step_cm * k))
        for j, opt in enumerate(row):
            if j > 0:
                p.add_run().add_tab()
            highlight = with_key and bool(opt.get("is_correct"))
            label_run = p.add_run(f"{opt['label']}. ")
            _set_font(label_run, bold=True, color=RED if highlight else None)
            # Chặn lỗi model nhân đôi ký tự đuôi phát âm ('catss' → 'cats') ngay khi
            # render để đề cũ tải lại là hết lỗi, không cần sinh lại (xem text_markup).
            option_text = dedupe_pronunciation_suffix(opt["text"])
            _add_runs_with_underline(p, option_text, bold=highlight, color=RED if highlight else None)


def build_exam_document(exam: Exam, variant: ExamVariant) -> Document:
    """Dựng đối tượng `Document` từ đề đã duyệt. Tách khỏi `render_exam_docx` để test
    được bố cục (thụt lề, đánh số, ẩn/hiện tiêu đề Phần con) mà không cần bọc
    StreamingResponse hay dựng cả DB."""
    with_key = exam.export_mode == ExportMode.ANSWER_KEY

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = section.right_margin = Cm(MARGIN_CM)

    student_line_p = _new_paragraph(doc)
    student_run = student_line_p.add_run("Full name: .......................................... Class: ..........")
    _set_font(student_run)

    _new_paragraph(doc)

    title_p = _new_paragraph(doc, center=True)
    title_run = title_p.add_run(f"{exam.title.upper()} — MÃ ĐỀ {variant.code}")
    _set_font(title_run, size=14, bold=True)

    ordered_blocks = sorted(exam.blocks, key=lambda b: b.order_no)
    question_no = 0
    for idx, block in enumerate(ordered_blocks):
        _new_paragraph(doc)  # dòng trống ngăn cách phần

        point_label = "point" if block.points == 1 else "points"
        heading_p = _new_paragraph(doc)
        heading_run = heading_p.add_run(
            f"{ROMAN[idx] if idx < len(ROMAN) else idx + 1}. {block.title.upper()} ({block.points} {point_label})"
        )
        _set_font(heading_run, bold=True)
        if block.instruction:
            block_instruction_p = _new_paragraph(doc)
            block_instruction_run = block_instruction_p.add_run(block.instruction)
            _set_font(block_instruction_run)

        order = variant.question_order.get(str(block.id), [str(q.id) for q in block.questions])
        by_id = {str(q.id): q for q in block.questions}
        ordered_questions = [by_id[qid] for qid in order if qid in by_id]
        parts_by_id = {p.id: p for p in block.parts}
        # Với block Pronunciation, các Phần con (Đuôi -s/-es, Đuôi -ed, Âm trong từ)
        # chỉ là công cụ chia kiểu câu khi sinh đề — học sinh không cần thấy tiêu đề
        # phần con, và cả khối đánh số liên tục 1..15 như một bài tập duy nhất (yêu
        # cầu chủ dự án 21/07/2026). Khác với các dạng như SENTENCE TRANSFORMATION
        # nơi phần con là nội dung đề thật, có tiêu đề riêng và đánh số lại từ 1.
        hide_parts = block.exercise_type.code == "pronunciation"
        shared_instruction_qids = _shared_instruction_question_ids(ordered_questions)

        current_part_id = None
        started = False
        # Câu dẫn ("Choose the word...") thụt vào trong; lựa chọn A/B/C/D nằm sát lề
        # trái cho block ẩn phần con — với phần con hiện tiêu đề thì cả hai cùng thụt.
        prompt_indent_cm = 0.0
        option_indent_cm = 0.0
        part = None
        part_question_no = 0
        part_letter_index = 0  # đếm phần con để gán nhãn A/B/C cho block ẩn phần con
        # Nhiều câu cùng khối (vd dạng phát âm/trọng âm) thường lặp lại y hệt 1 câu
        # dẫn/hướng dẫn — chỉ in 1 lần cho lần đầu xuất hiện, các câu sau chỉ còn số
        # thứ tự + lựa chọn (đúng format đề tham khảo giáo viên gửi). Reset khi sang
        # block/phần mới vì mỗi block/phần có thể có câu dẫn khác nhau.
        last_prompt_text: str | None = None
        for question in ordered_questions:
            if not started or question.part_id != current_part_id:
                started = True
                current_part_id = question.part_id
                last_prompt_text = None
                part_question_no = 0
                part = parts_by_id.get(current_part_id) if current_part_id else None
                # Câu dẫn/tiêu đề phần con luôn thụt vào trong; lựa chọn thụt theo khi
                # phần con hiện tiêu đề (phân cấp I. > 1. > câu hỏi), nhưng nằm sát lề
                # trái khi ẩn phần con (Pronunciation) để chỉ dòng hướng dẫn mới thụt.
                if part is not None:
                    prompt_indent_cm = PART_CONTENT_INDENT_CM
                    option_indent_cm = 0.0 if hide_parts else PART_CONTENT_INDENT_CM
                else:
                    prompt_indent_cm = option_indent_cm = 0.0
                if part is not None and hide_parts:
                    # Phần con Pronunciation: in 1 dòng hướng dẫn đánh nhãn chữ cái
                    # A/B/C (thay cho cả tiêu đề phần con lẫn câu dẫn lặp theo từng
                    # câu) = câu dẫn chung của phần, thụt vào trong. Mọi câu bên dưới
                    # chỉ còn số thứ tự (1..15 liên tục) + lựa chọn nằm sát lề trái.
                    part_letter_index += 1
                    letter = chr(ord("A") + part_letter_index - 1)
                    instruction_text = question.prompt_text or part.title
                    part_line_p = _new_paragraph(doc)
                    part_line_p.paragraph_format.left_indent = Cm(prompt_indent_cm)
                    letter_run = part_line_p.add_run(f"{letter}. ")
                    _set_font(letter_run, bold=True)
                    _add_runs_with_underline(part_line_p, instruction_text)
                    # Chặn in lại câu dẫn theo từng câu (đã gộp vào dòng hướng dẫn trên).
                    last_prompt_text = question.prompt_text
                else:
                    if part is not None:
                        part_heading_p = _new_paragraph(doc)
                        part_heading_p.paragraph_format.left_indent = Cm(prompt_indent_cm)
                        part_heading_run = part_heading_p.add_run(f"{part.order_no}. {part.title}")
                        _set_font(part_heading_run, bold=True)
                        if part.instruction:
                            part_instruction_p = _new_paragraph(doc)
                            part_instruction_p.paragraph_format.left_indent = Cm(prompt_indent_cm)
                            part_instruction_run = part_instruction_p.add_run(part.instruction)
                            _set_font(part_instruction_run)
                    if question.id in shared_instruction_qids:
                        # Nhóm dùng chung 1 câu dẫn (vd khối STRESS không có phần con):
                        # in câu dẫn thành dòng riêng, KHÔNG gộp số thứ tự vào — nhờ vậy
                        # câu đầu nhóm vẫn giữ số ở dòng lựa chọn như các câu còn lại.
                        instruction_p = _new_paragraph(doc)
                        instruction_p.paragraph_format.left_indent = Cm(prompt_indent_cm)
                        _add_runs_with_underline(instruction_p, question.prompt_text)
                        last_prompt_text = question.prompt_text

            question_no += 1
            # Phần con hiện tiêu đề đánh số lại từ 1 riêng cho phần đó (mỗi phần con là
            # 1 bài tập độc lập). Block ẩn phần con (Pronunciation) và câu không thuộc
            # phần nào đánh số liên tục xuyên suốt khối (1..15).
            if part is not None and not hide_parts:
                part_question_no += 1
                display_no = part_question_no
            else:
                display_no = question_no

            if question.passage_text:
                passage_p = _new_paragraph(doc, justify=True)
                passage_p.paragraph_format.left_indent = Cm(prompt_indent_cm)
                _add_runs_with_underline(passage_p, question.passage_text)

            # Block ẩn phần con (Pronunciation) đã in câu dẫn chung ở dòng nhãn A/B/C —
            # mọi câu chỉ còn số thứ tự + lựa chọn, không in lại câu dẫn theo từng câu.
            # Các dạng khác: câu dẫn giống câu trước (cùng phần) chỉ in 1 lần, câu sau
            # in gộp số thứ tự cùng dòng lựa chọn đầu tiên (báo cáo giáo viên 21/07/2026).
            suppress_inline_prompt = (hide_parts and part is not None) or question.id in shared_instruction_qids
            show_prompt_text = (
                not suppress_inline_prompt
                and bool(question.prompt_text)
                and question.prompt_text != last_prompt_text
            )
            if show_prompt_text:
                prompt_p = _new_paragraph(doc)
                prompt_p.paragraph_format.left_indent = Cm(prompt_indent_cm)
                no_run = prompt_p.add_run(f"{display_no}.")
                _set_font(no_run, bold=True)
                prompt_p.add_run(" ")
                _add_runs_with_underline(prompt_p, question.prompt_text)
                last_prompt_text = question.prompt_text

            if question.options:
                _render_options(
                    doc,
                    question.options,
                    with_key,
                    indent_cm=option_indent_cm,
                    number=None if show_prompt_text else f"{display_no}.",
                )
            elif with_key:
                answer_p = _new_paragraph(doc)
                answer_p.paragraph_format.left_indent = Cm(option_indent_cm)
                prefix = "" if show_prompt_text else f"{display_no}. "
                answer_run = answer_p.add_run(f"{prefix}Đáp án: {question.answer_text}")
                _set_font(answer_run, bold=True, color=RED)
            elif not show_prompt_text:
                # Câu không có lựa chọn, không hiện đáp án (không with_key), và không
                # có câu dẫn mới — vẫn phải in số thứ tự để không mất dấu vết câu hỏi.
                fallback_p = _new_paragraph(doc)
                fallback_p.paragraph_format.left_indent = Cm(option_indent_cm)
                no_run = fallback_p.add_run(f"{display_no}.")
                _set_font(no_run, bold=True)
    return doc


def render_exam_docx(exam: Exam, variant: ExamVariant) -> StreamingResponse:
    doc = build_exam_document(exam, variant)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    with_key = exam.export_mode == ExportMode.ANSWER_KEY
    suffix = "dap-an-do" if with_key else "hoc-sinh"
    filename = f"de-ma-{variant.code.lower()}-{suffix}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
