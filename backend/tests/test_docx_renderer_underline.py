"""Test riêng cho phần render <u>...</u> (dạng phát âm/trọng âm cần gạch chân, xem
app/services/prompts.py) — không dựng cả Exam/ExamVariant, chỉ test helper thuần."""

from docx import Document

from app.services.docx_renderer import _add_runs_with_underline, _visible_length


def _paragraph_runs(text: str, **kwargs):
    doc = Document()
    p = doc.add_paragraph()
    _add_runs_with_underline(p, text, **kwargs)
    return p.runs


def test_plain_text_without_markup_produces_single_run():
    runs = _paragraph_runs("bread")
    assert len(runs) == 1
    assert runs[0].text == "bread"
    assert runs[0].font.underline is not True


def test_markup_in_middle_splits_into_three_runs():
    runs = _paragraph_runs("br<u>ea</u>d")
    texts = [r.text for r in runs]
    assert texts == ["br", "ea", "d"]
    assert runs[0].font.underline is not True
    assert runs[1].font.underline is True
    assert runs[2].font.underline is not True


def test_underlined_part_is_always_bold_even_without_highlight():
    """Phần gạch chân luôn in đậm để dễ nhận ra, kể cả khi lựa chọn không phải đáp
    án đúng (bold=False mặc định) — yêu cầu chủ dự án 21/07/2026."""
    runs = _paragraph_runs("br<u>ea</u>d")
    texts = [r.text for r in runs]
    assert runs[texts.index("ea")].font.bold is True
    assert runs[texts.index("br")].font.bold is not True
    assert runs[texts.index("d")].font.bold is not True


def test_markup_at_start_has_no_leading_empty_run():
    runs = _paragraph_runs("<u>han</u>dsome")
    texts = [r.text for r in runs]
    assert texts == ["han", "dsome"]
    assert runs[0].font.underline is True


def test_markup_at_end_has_no_trailing_empty_run():
    runs = _paragraph_runs("be<u>gin</u>")
    texts = [r.text for r in runs]
    assert texts == ["be", "gin"]
    assert runs[1].font.underline is True


def test_bold_and_color_apply_to_all_runs_when_highlighted():
    from docx.shared import RGBColor

    red = RGBColor(0xC0, 0x00, 0x00)
    runs = _paragraph_runs("br<u>ea</u>d", bold=True, color=red)
    for run in runs:
        assert run.font.bold is True
        assert run.font.color.rgb == red


def test_visible_length_excludes_markup_characters():
    assert _visible_length("bread") == 5
    assert _visible_length("br<u>ea</u>d") == 5
    assert _visible_length("plain text option") == len("plain text option")
