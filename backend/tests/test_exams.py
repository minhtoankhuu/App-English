import io
import uuid

import pytest
from docx import Document
from sqlalchemy import select

from app.models.academic import Grade, ProficiencyLevel, Unit
from app.models.exercise import ExerciseType
from app.models.grammar import GrammarGroup, GrammarPoint, GrammarTopic
from app.models.user import User, UserRole
from app.security import hash_password
from app.services.usage import reserve_usage


def _login_as_teacher(client, db):
    user = User(
        email="teacher@examcraft.dev",
        password_hash=hash_password("Secret123!"),
        full_name="Teacher",
        role=UserRole.TEACHER,
    )
    db.add(user)
    db.commit()
    response = client.post("/auth/login", json={"email": "teacher@examcraft.dev", "password": "Secret123!"})
    assert response.status_code == 200
    return user


def _login_as_admin(client, db):
    user = User(
        email="exam-admin@examcraft.dev",
        password_hash=hash_password("Secret123!"),
        full_name="Exam Admin",
        role=UserRole.ADMIN,
    )
    db.add(user)
    db.commit()
    response = client.post("/auth/login", json={"email": user.email, "password": "Secret123!"})
    assert response.status_code == 200
    return user


def _grade7_unit3(db):
    grade7 = db.scalar(select(Grade).where(Grade.number == 7))
    unit3 = db.scalar(select(Unit).where(Unit.grade_id == grade7.id, Unit.order_no == 3))
    assert unit3.title == "Community Service"
    return grade7, unit3


def _level(db, code):
    return db.scalar(select(ProficiencyLevel).where(ProficiencyLevel.code == code))


def _exercise_type(db, code):
    return db.scalar(select(ExerciseType).where(ExerciseType.code == code))


def _create_golden_exam(client, db):
    grade7, unit3 = _grade7_unit3(db)
    level_a2 = _level(db, "A2")
    resp = client.post(
        "/exams",
        json={
            "title": "Unit 3 — Revision Test",
            "grade_id": str(grade7.id),
            "level_id": str(level_a2.id),
            "source_type": "global_success",
            "unit_id": str(unit3.id),
        },
    )
    assert resp.status_code == 201
    return resp.json()


@pytest.mark.parametrize(
    ("method", "path", "payload"),
    [
        ("get", "/exams", None),
        ("post", "/exams", {}),
        ("get", f"/exams/{uuid.uuid4()}", None),
    ],
)
def test_admin_cannot_access_exam_workflow(client, db, method, path, payload):
    _login_as_admin(client, db)

    request = getattr(client, method)
    response = request(path, json=payload) if payload is not None else request(path)

    assert response.status_code == 403
    assert response.json()["detail"] == "Không đủ quyền truy cập"


def test_unauthenticated_cannot_access_exam_workflow(client):
    response = client.get("/exams")

    assert response.status_code == 401


def test_create_exam_rejects_mismatched_source_fields(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    grade7, unit3 = _grade7_unit3(seeded_db)
    level_a2 = _level(seeded_db, "A2")

    response = client.post(
        "/exams",
        json={
            "title": "Bad exam",
            "grade_id": str(grade7.id),
            "level_id": str(level_a2.id),
            "source_type": "global_success",
            # thiếu unit_id
        },
    )
    assert response.status_code == 400


def test_other_teacher_cannot_access_exam(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)

    other = User(
        email="other@examcraft.dev",
        password_hash=hash_password("Secret123!"),
        full_name="Other",
        role=UserRole.TEACHER,
    )
    seeded_db.add(other)
    seeded_db.commit()
    client.post("/auth/logout")
    client.post("/auth/login", json={"email": "other@examcraft.dev", "password": "Secret123!"})

    response = client.get(f"/exams/{exam['id']}")
    assert response.status_code == 403


def test_full_golden_flow_create_generate_review_export(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)
    exam_id = exam["id"]

    block_specs = [
        ("pronunciation", "I. Pronunciation", 2, "1.0"),
        ("multiple_choice", "II. Vocabulary & Grammar", 2, "3.0"),
        ("reading_true_false", "III. Reading", 1, "1.0"),
        ("sentence_rewrite", "IV. Sentence Transformation", 1, "3.0"),
    ]
    block_ids = []
    for code, title, count, points in block_specs:
        ex_type = _exercise_type(seeded_db, code)
        resp = client.post(
            f"/exams/{exam_id}/blocks",
            json={
                "exercise_type_id": str(ex_type.id),
                "title": title,
                "question_count": count,
                "points": points,
            },
        )
        assert resp.status_code == 201, resp.text
        block_ids.append(resp.json()["id"])

    # sinh câu hỏi
    resp = client.post(f"/exams/{exam_id}/generate")
    assert resp.status_code == 200
    assert client.get("/usage/me").json()["used"] == 4
    detail = resp.json()
    assert detail["status"] == "draft"
    total_questions = sum(len(b["questions"]) for b in detail["blocks"])
    assert total_questions == 6

    # câu bị động (IV) phải cảnh báo vượt trình độ (B1 > A2 của đề)
    rewrite_block = next(b for b in detail["blocks"] if b["exercise_type"]["code"] == "sentence_rewrite")
    rewrite_q = rewrite_block["questions"][0]
    assert any("vượt trình độ" in w for w in rewrite_q["warnings"])
    assert rewrite_q["answer_text"] == "will be built"

    # câu phát âm phải đúng nội dung golden fixture
    pron_block = next(b for b in detail["blocks"] if b["exercise_type"]["code"] == "pronunciation")
    assert pron_block["questions"][0]["answer_text"] == "B. bread"

    # chưa duyệt hết thì complete-review phải 409
    resp = client.post(f"/exams/{exam_id}/complete-review")
    assert resp.status_code == 409

    # duyệt toàn bộ câu
    all_question_ids = [q["id"] for b in detail["blocks"] for q in b["questions"]]
    for qid in all_question_ids:
        resp = client.patch(f"/exams/{exam_id}/questions/{qid}", json={"is_approved": True})
        assert resp.status_code == 200
        assert resp.json()["is_approved"] is True

    resp = client.post(f"/exams/{exam_id}/complete-review")
    assert resp.status_code == 200
    assert resp.json()["status"] == "reviewed"

    # lưu cấu hình xuất -> Sẵn sàng xuất, tạo 2 mã đề
    resp = client.post(
        f"/exams/{exam_id}/export-config", json={"export_mode": "answer_key", "variant_count": 2}
    )
    assert resp.status_code == 200
    ready = resp.json()
    assert ready["status"] == "ready"
    assert ready["variant_count"] == 2

    # tải DOCX mã đề A, kiểm tra mở được và có nội dung đúng
    resp = client.get(f"/exams/{exam_id}/export.docx", params={"variant": "A"})
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    doc = Document(io.BytesIO(resp.content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "UNIT 3 — REVISION TEST" in full_text
    assert "MÃ ĐỀ A" in full_text
    assert "PRONUNCIATION" in full_text
    assert "School:" in full_text
    assert "Full name:" in full_text
    assert "Mark:" in full_text
    assert "Time: 45 minutes" in full_text
    # đáp án tô đỏ: chữ "bread" xuất hiện, và có ít nhất 1 run màu đỏ trong file
    red_runs = [r for p in doc.paragraphs for r in p.runs if r.font.color and r.font.color.rgb is not None]
    assert len(red_runs) > 0

    # mã đề không tồn tại (chỉ tạo A, B) -> 404
    resp = client.get(f"/exams/{exam_id}/export.docx", params={"variant": "C"})
    assert resp.status_code == 404


def test_regenerate_blocked_when_locked_or_approved(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)
    ex_type = _exercise_type(seeded_db, "multiple_choice")
    resp = client.post(
        f"/exams/{exam['id']}/blocks",
        json={"exercise_type_id": str(ex_type.id), "title": "II. Grammar", "question_count": 1, "points": "2.0"},
    )
    block_id = resp.json()["id"]
    client.post(f"/exams/{exam['id']}/generate")
    detail = client.get(f"/exams/{exam['id']}").json()
    question_id = detail["blocks"][0]["questions"][0]["id"]

    # khóa câu -> sinh lại phải 409
    resp = client.patch(f"/exams/{exam['id']}/questions/{question_id}", json={"is_locked": True})
    assert resp.status_code == 200
    resp = client.post(f"/exams/{exam['id']}/questions/{question_id}/regenerate")
    assert resp.status_code == 409

    # mở khóa, duyệt -> sinh lại cũng phải 409
    client.patch(f"/exams/{exam['id']}/questions/{question_id}", json={"is_locked": False})
    client.patch(f"/exams/{exam['id']}/questions/{question_id}", json={"is_approved": True})
    resp = client.post(f"/exams/{exam['id']}/questions/{question_id}/regenerate")
    assert resp.status_code == 409

    # bỏ duyệt -> sinh lại được
    client.patch(f"/exams/{exam['id']}/questions/{question_id}", json={"is_approved": False})
    resp = client.post(f"/exams/{exam['id']}/questions/{question_id}/regenerate")
    assert resp.status_code == 200
    assert resp.json()["is_approved"] is False
    assert client.get("/usage/me").json()["used"] == 2
    _ = block_id


def test_generate_rejects_whole_exam_when_remaining_is_insufficient(client, seeded_db):
    teacher = _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)
    for code, title in [("pronunciation", "I"), ("multiple_choice", "II")]:
        exercise_type = _exercise_type(seeded_db, code)
        response = client.post(
            f"/exams/{exam['id']}/blocks",
            json={"exercise_type_id": str(exercise_type.id), "title": title, "question_count": 1, "points": "1.0"},
        )
        assert response.status_code == 201
    reserve_usage(seeded_db, teacher, 9)
    seeded_db.commit()

    response = client.post(f"/exams/{exam['id']}/generate")

    assert response.status_code == 429
    assert response.json()["detail"]["remaining"] == 1
    assert client.get(f"/exams/{exam['id']}").json()["blocks"][0]["questions"] == []
    assert client.get("/usage/me").json()["used"] == 9


def test_reorder_blocks(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)
    type_a = _exercise_type(seeded_db, "pronunciation")
    type_b = _exercise_type(seeded_db, "multiple_choice")

    b1 = client.post(
        f"/exams/{exam['id']}/blocks",
        json={"exercise_type_id": str(type_a.id), "title": "I", "question_count": 1, "points": "1.0"},
    ).json()
    b2 = client.post(
        f"/exams/{exam['id']}/blocks",
        json={"exercise_type_id": str(type_b.id), "title": "II", "question_count": 1, "points": "1.0"},
    ).json()
    assert b1["order_no"] == 1
    assert b2["order_no"] == 2

    resp = client.post(f"/exams/{exam['id']}/blocks/reorder", json={"block_ids": [b2["id"], b1["id"]]})
    assert resp.status_code == 200
    blocks = resp.json()["blocks"]
    ordered = sorted(blocks, key=lambda b: b["order_no"])
    assert ordered[0]["id"] == b2["id"]
    assert ordered[1]["id"] == b1["id"]


def test_export_before_review_returns_409(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)
    resp = client.post(
        f"/exams/{exam['id']}/export-config", json={"export_mode": "plain", "variant_count": 1}
    )
    assert resp.status_code == 409

    resp = client.get(f"/exams/{exam['id']}/export.docx")
    assert resp.status_code == 409


def test_update_exam_validates_source_fields_even_without_source_type_in_payload(client, seeded_db):
    """Regression: PATCH chỉ đổi unit_id (không kèm source_type) vẫn phải được kiểm tra
    nhất quán với source_type hiện tại của đề — trước đây validation bị bỏ qua trong
    trường hợp này."""
    _login_as_teacher(client, seeded_db)
    exam = _create_golden_exam(client, seeded_db)  # global_success, có unit_id

    # đổi unit_id sang Unit khác cùng lớp -> hợp lệ (vẫn global_success, vẫn có unit_id)
    grade7, _ = _grade7_unit3(seeded_db)
    unit2 = seeded_db.scalar(select(Unit).where(Unit.grade_id == grade7.id, Unit.order_no == 2))
    resp = client.patch(f"/exams/{exam['id']}", json={"unit_id": str(unit2.id)})
    assert resp.status_code == 200
    assert resp.json()["unit_id"] == str(unit2.id)

    # đặt cambridge_certificate_id trong khi đề vẫn là global_success -> phải bị từ chối
    certs = client.get("/catalog/cambridge-certificates").json()
    resp = client.patch(
        f"/exams/{exam['id']}", json={"cambridge_certificate_id": certs[0]["id"]}
    )
    assert resp.status_code == 400


def test_common_knowledge_exam_grammar_selection(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    grade7, _ = _grade7_unit3(seeded_db)
    level_a2 = _level(seeded_db, "A2")
    topic = seeded_db.scalar(select(GrammarTopic).where(GrammarTopic.code == "tense"))
    points = seeded_db.scalars(
        select(GrammarPoint)
        .join(GrammarGroup, GrammarPoint.group_id == GrammarGroup.id)
        .where(GrammarGroup.topic_id == topic.id)
    ).all()
    point_ids = [str(p.id) for p in points[:2]]

    resp = client.post(
        "/exams",
        json={
            "title": "Tense revision",
            "grade_id": str(grade7.id),
            "level_id": str(level_a2.id),
            "source_type": "common_knowledge",
            "grammar_topic_id": str(topic.id),
        },
    )
    assert resp.status_code == 201
    exam = resp.json()
    assert exam["grammar_point_ids"] == []

    resp = client.put(f"/exams/{exam['id']}/grammar-selection", json={"grammar_point_ids": point_ids})
    assert resp.status_code == 200
    assert sorted(resp.json()["grammar_point_ids"]) == sorted(point_ids)

    # gọi lại với danh sách khác phải THAY THẾ, không cộng dồn
    resp = client.put(f"/exams/{exam['id']}/grammar-selection", json={"grammar_point_ids": [point_ids[0]]})
    assert resp.status_code == 200
    assert resp.json()["grammar_point_ids"] == [point_ids[0]]


def test_list_exams_returns_only_own_exams_with_summary(client, seeded_db):
    _login_as_teacher(client, seeded_db)
    _create_golden_exam(client, seeded_db)
    _create_golden_exam(client, seeded_db)

    resp = client.get("/exams")
    assert resp.status_code == 200
    exams = resp.json()
    assert len(exams) == 2
    for e in exams:
        assert e["status"] == "draft"
        assert e["grade_number"] == 7
        assert e["level_code"] == "A2"
        assert e["total_questions"] == 0
