from pathlib import Path

from docx import Document

from app.models.knowledge import DocumentChunkType
from app.services.knowledge_parser import parse_lesson_docx

KB_ROOT = Path(__file__).resolve().parents[2] / "Knowledge_Base" / "Global Success"


def test_parses_golden_reference_unit3_grade7():
    chunks = parse_lesson_docx(KB_ROOT / "G7" / "GS7 - UNIT 3 - LESSON.docx")

    types_present = {c.chunk_type for c in chunks}
    assert types_present == {
        DocumentChunkType.VOCABULARY,
        DocumentChunkType.WORD_FORM,
        DocumentChunkType.PHRASE,
        DocumentChunkType.GRAMMAR,
    }

    volunteer = next(
        c for c in chunks if c.chunk_type == DocumentChunkType.VOCABULARY and c.structured and c.structured["word"] == "volunteer"
    )
    assert volunteer.structured == {
        "word": "volunteer",
        "ipa": "ˌvɑːlənˈtɪr",
        "pos": "n, v",
        "meaning": "tình nguyện viên; tình nguyện",
    }

    grammar_table_chunks = [c for c in chunks if c.chunk_type == DocumentChunkType.GRAMMAR and "Simple Past Tense" in c.raw_text]
    assert grammar_table_chunks, "bảng ngữ pháp (Simple Past Tense) phải được trích từ table"
    table_chunk = grammar_table_chunks[0]
    assert table_chunk.structured is not None
    grid = table_chunk.structured["table"]
    assert isinstance(grid, list) and isinstance(grid[0], list)
    assert any("Simple Past Tense" in cell for row in grid for cell in row)

    order_nos = [c.order_no for c in chunks]
    assert order_nos == sorted(order_nos)
    assert len(order_nos) == len(set(order_nos))


def _all_lesson_files() -> list[Path]:
    return sorted(f for grade in ("G6", "G7", "G8") for f in (KB_ROOT / grade).glob("*.docx"))


def test_vocabulary_match_rate_across_all_units_is_high():
    all_files = _all_lesson_files()
    assert len(all_files) == 36

    total_vocab = 0
    total_matched = 0
    for file_path in all_files:
        chunks = parse_lesson_docx(file_path)
        vocab_chunks = [c for c in chunks if c.chunk_type == DocumentChunkType.VOCABULARY]
        total_vocab += len(vocab_chunks)
        total_matched += sum(1 for c in vocab_chunks if c.structured is not None)

    assert total_vocab > 2000
    assert total_matched / total_vocab > 0.9


def test_every_unit_file_has_a_minimum_vocabulary_count():
    """Một số Unit dùng tiêu đề "New words"/"Vocabulary" in đậm nhưng không viết HOA
    (khác với "VOCABULARY" viết HOA của đa số Unit) — nếu parser chỉ nhận diện tiêu đề
    bằng viết HOA, cả file sẽ không có mục từ vựng nào được phân loại đúng (rơi hết vào
    DocumentChunkType.OTHER). Ngưỡng thấp (5) đủ để bắt lỗi này ở TỪNG file thay vì chỉ
    nhìn tổng toàn kho (tổng vẫn > 2000 dù một vài file bằng 0, xem test phía trên)."""
    for file_path in _all_lesson_files():
        chunks = parse_lesson_docx(file_path)
        vocab_count = sum(1 for c in chunks if c.chunk_type == DocumentChunkType.VOCABULARY)
        assert vocab_count >= 5, f"{file_path.name} chỉ có {vocab_count} mục từ vựng — tiêu đề có thể không được nhận diện"


def _build_bold_title_case_header_doc(tmp_path) -> Path:
    doc = Document()
    doc.add_paragraph("UNIT 99: TEST UNIT")

    header = doc.add_paragraph()
    header.add_run("New words").bold = True
    doc.add_paragraph("ability /əˈbɪləti/ (n) : khả năng")
    doc.add_paragraph("amazing /əˈmeɪzɪŋ/ (adj) : đáng kinh ngạc")

    header2 = doc.add_paragraph()
    header2.add_run("Grammar and Structures").bold = True
    doc.add_paragraph("She has lived here since 2010.")

    path = tmp_path / "bold-title-case-headers.docx"
    doc.save(str(path))
    return path


def test_recognizes_bold_title_case_headers_not_just_all_caps(tmp_path):
    path = _build_bold_title_case_header_doc(tmp_path)

    chunks = parse_lesson_docx(path)

    vocab_chunks = [c for c in chunks if c.chunk_type == DocumentChunkType.VOCABULARY]
    assert len(vocab_chunks) == 2
    assert vocab_chunks[0].section_title == "New words"
    assert vocab_chunks[0].structured == {"word": "ability", "ipa": "əˈbɪləti", "pos": "n", "meaning": "khả năng"}

    grammar_chunks = [c for c in chunks if c.chunk_type == DocumentChunkType.GRAMMAR]
    assert len(grammar_chunks) == 1
    assert grammar_chunks[0].section_title == "Grammar and Structures"


def _build_numbered_word_form_doc(tmp_path) -> Path:
    doc = Document()
    doc.add_paragraph("UNIT 99: TEST UNIT")
    header = doc.add_paragraph()
    header.add_run("WORD FORM").bold = True
    doc.add_paragraph("7. responsible (adj): có trách nhiệm")
    doc.add_paragraph("10. harm (v/n): gây hại / sự tổn hại\n11. effective (adj): hiệu quả")
    path = tmp_path / "numbered-word-form.docx"
    doc.save(str(path))
    return path


def test_strips_manual_list_numbering_from_word_form_entries(tmp_path):
    """Một số Unit đánh số thủ công cho danh sách WORD FORM ("7. responsible (adj): ...",
    có khi 2 mục gộp trong 1 đoạn văn nối bằng xuống dòng mềm "...\n11. effective..."). Số
    này chỉ là marker liệt kê, không phải nội dung — phải bị bỏ khỏi raw_text."""
    path = _build_numbered_word_form_doc(tmp_path)

    chunks = parse_lesson_docx(path)
    word_form_chunks = [c for c in chunks if c.chunk_type == DocumentChunkType.WORD_FORM]

    assert word_form_chunks[0].raw_text == "responsible (adj): có trách nhiệm"
    assert word_form_chunks[1].raw_text == "harm (v/n): gây hại / sự tổn hại\neffective (adj): hiệu quả"


def test_parser_never_raises_on_any_unit_file():
    for file_path in _all_lesson_files():
        chunks = parse_lesson_docx(file_path)
        assert len(chunks) > 0
        assert all(c.raw_text.strip() for c in chunks)
