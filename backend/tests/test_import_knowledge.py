import shutil
from pathlib import Path

from docx import Document
from sqlalchemy import func, select

from app.import_knowledge import import_global_success, import_grade9_vocabulary
from app.models.academic import Grade, Unit
from app.models.knowledge import KnowledgeChunk, KnowledgeDocument

KB_ROOT = Path(__file__).resolve().parents[2] / "Knowledge_Base" / "Global Success"
_G9_VOCAB_FILE = "Vocabulary Global Success 9 (1).docx"


def _copy_unit3_grade7(tmp_path: Path) -> Path:
    dest_dir = tmp_path / "Global Success" / "G7"
    dest_dir.mkdir(parents=True)
    shutil.copyfile(KB_ROOT / "G7" / "GS7 - UNIT 3 - LESSON.docx", dest_dir / "GS7 - UNIT 3 - LESSON.docx")
    return tmp_path


def _unit3_grade7(db) -> Unit:
    return db.scalar(select(Unit).join(Grade).where(Grade.number == 7, Unit.order_no == 3))


def _chunk_count(db, document_id) -> int:
    return db.scalar(select(func.count()).select_from(KnowledgeChunk).where(KnowledgeChunk.document_id == document_id))


def test_import_creates_document_and_chunks(seeded_db, tmp_path):
    base_path = _copy_unit3_grade7(tmp_path)

    stats = import_global_success(seeded_db, base_path)
    seeded_db.commit()

    assert stats.files_seen == 1
    assert stats.documents_created == 1
    assert stats.chunks_written > 0

    unit3 = _unit3_grade7(seeded_db)
    document = seeded_db.scalar(select(KnowledgeDocument).where(KnowledgeDocument.unit_id == unit3.id))
    assert document is not None
    assert document.file_name == "GS7 - UNIT 3 - LESSON.docx"
    assert _chunk_count(seeded_db, document.id) == stats.chunks_written


def test_import_is_idempotent_on_rerun(seeded_db, tmp_path):
    base_path = _copy_unit3_grade7(tmp_path)

    import_global_success(seeded_db, base_path)
    seeded_db.commit()
    unit3 = _unit3_grade7(seeded_db)
    document = seeded_db.scalar(select(KnowledgeDocument).where(KnowledgeDocument.unit_id == unit3.id))
    rows_after_first_import = _chunk_count(seeded_db, document.id)

    stats = import_global_success(seeded_db, base_path)
    seeded_db.commit()

    assert stats.documents_created == 0
    assert stats.documents_updated == 0
    assert stats.documents_unchanged == 1
    assert _chunk_count(seeded_db, document.id) == rows_after_first_import


def test_import_replaces_chunks_when_file_content_changes(seeded_db, tmp_path):
    base_path = _copy_unit3_grade7(tmp_path)
    dest_file = base_path / "Global Success" / "G7" / "GS7 - UNIT 3 - LESSON.docx"

    import_global_success(seeded_db, base_path)
    seeded_db.commit()
    unit3 = _unit3_grade7(seeded_db)
    document = seeded_db.scalar(select(KnowledgeDocument).where(KnowledgeDocument.unit_id == unit3.id))
    original_checksum = document.checksum

    document_xml = Document(str(dest_file))  # giả lập file được cập nhật: sửa 1 đoạn rồi lưu lại
    document_xml.paragraphs[0].add_run(" (updated)")
    document_xml.save(str(dest_file))

    stats = import_global_success(seeded_db, base_path)
    seeded_db.commit()
    seeded_db.refresh(document)

    assert stats.documents_updated == 1
    assert document.checksum != original_checksum
    assert _chunk_count(seeded_db, document.id) > 0


def test_import_force_reparses_even_when_checksum_unchanged(seeded_db, tmp_path):
    base_path = _copy_unit3_grade7(tmp_path)

    import_global_success(seeded_db, base_path)
    seeded_db.commit()
    unit3 = _unit3_grade7(seeded_db)
    document = seeded_db.scalar(select(KnowledgeDocument).where(KnowledgeDocument.unit_id == unit3.id))
    original_checksum = document.checksum

    stats = import_global_success(seeded_db, base_path, force=True)
    seeded_db.commit()
    seeded_db.refresh(document)

    assert stats.documents_updated == 1
    assert stats.documents_unchanged == 0
    assert document.checksum == original_checksum
    assert _chunk_count(seeded_db, document.id) == stats.chunks_written


def _copy_g9_vocab(tmp_path: Path) -> Path:
    dest_dir = tmp_path / "Global Success" / "G9"
    dest_dir.mkdir(parents=True)
    shutil.copyfile(KB_ROOT / "G9" / _G9_VOCAB_FILE, dest_dir / _G9_VOCAB_FILE)
    return tmp_path


def _grade9_units(db) -> list[Unit]:
    grade9 = db.scalar(select(Grade).where(Grade.number == 9))
    return list(db.scalars(select(Unit).where(Unit.grade_id == grade9.id)))


def test_import_g9_vocabulary_creates_per_unit_docs(seeded_db, tmp_path):
    assert _grade9_units(seeded_db), "seed phải có Unit lớp 9"
    base_path = _copy_g9_vocab(tmp_path)

    stats = import_grade9_vocabulary(seeded_db, base_path)
    seeded_db.commit()

    assert stats.documents_created >= 1
    assert stats.chunks_written > 0
    # mỗi Unit lớp 9 có tài liệu từ vựng với chunk VOCABULARY
    unit1 = seeded_db.scalar(
        select(Unit).join(Grade).where(Grade.number == 9, Unit.order_no == 1)
    )
    doc = seeded_db.scalar(select(KnowledgeDocument).where(KnowledgeDocument.unit_id == unit1.id))
    assert doc is not None and doc.file_name == _G9_VOCAB_FILE
    assert _chunk_count(seeded_db, doc.id) > 0


def test_import_g9_vocabulary_is_idempotent(seeded_db, tmp_path):
    base_path = _copy_g9_vocab(tmp_path)
    import_grade9_vocabulary(seeded_db, base_path)
    seeded_db.commit()

    stats = import_grade9_vocabulary(seeded_db, base_path)
    seeded_db.commit()

    assert stats.documents_created == 0
    assert stats.documents_updated == 0
    assert stats.documents_unchanged >= 1
