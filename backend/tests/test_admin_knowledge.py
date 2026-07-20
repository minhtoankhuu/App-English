from pathlib import Path

from docx import Document
from sqlalchemy import select

from app.models.academic import Grade, Unit
from app.models.grammar import GrammarPoint
from app.models.knowledge import KnowledgeDocument
from app.models.user import User, UserRole
from app.security import hash_password

KB_FILE = (
    Path(__file__).resolve().parents[2]
    / "Knowledge_Base"
    / "Global Success"
    / "G7"
    / "GS7 - UNIT 3 - LESSON.docx"
)


def _login(client, db, *, email: str = "kb-admin@examcraft.dev", role: UserRole = UserRole.ADMIN) -> User:
    user = User(email=email, password_hash=hash_password("Secret123!"), full_name="Admin User", role=role)
    db.add(user)
    db.commit()
    response = client.post("/auth/login", json={"email": email, "password": "Secret123!"})
    assert response.status_code == 200
    return user


def _unit3_grade7(db) -> Unit:
    return db.scalar(select(Unit).join(Grade).where(Grade.number == 7, Unit.order_no == 3))


def _unit2_grade7(db) -> Unit:
    return db.scalar(select(Unit).join(Grade).where(Grade.number == 7, Unit.order_no == 2))


def _present_simple_point(db) -> GrammarPoint:
    return db.scalar(select(GrammarPoint).where(GrammarPoint.name == "Present Simple"))


def _bold_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _build_grammar_reference_docx(tmp_path) -> bytes:
    doc = Document()
    _bold_paragraph(doc, "Cách dùng:")
    doc.add_paragraph("The earth is round. It goes around the sun.")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "V (+)"
    table.rows[0].cells[1].text = "S + V…"
    path = tmp_path / "present-simple.docx"
    doc.save(str(path))
    return path.read_bytes()


def _upload(client, unit_id, filename="GS7 - UNIT 3 - LESSON.docx", content: bytes | None = None):
    content = content if content is not None else KB_FILE.read_bytes()
    return client.post(
        "/admin/knowledge-documents",
        data={"unit_id": str(unit_id)},
        files={"file": (filename, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )


def _upload_grammar(client, grammar_point_id, filename="present-simple.docx", content: bytes = b""):
    return client.post(
        "/admin/knowledge-documents",
        data={"grammar_point_id": str(grammar_point_id)},
        files={"file": (filename, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )


def test_teacher_cannot_access_admin_knowledge(client, seeded_db):
    _login(client, seeded_db, email="teacher@examcraft.dev", role=UserRole.TEACHER)

    assert client.get("/admin/knowledge-documents").status_code == 403


def test_unauthenticated_cannot_access(client):
    assert client.get("/admin/knowledge-documents").status_code == 401


def test_upload_creates_document_and_chunks(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)

    resp = _upload(client, unit3.id)

    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["file_name"] == "GS7 - UNIT 3 - LESSON.docx"
    assert body["chunk_count"] > 0
    assert body["is_published"] is True
    assert body["unit"]["order_no"] == 3
    assert body["unit"]["grade_number"] == 7


def test_upload_same_checksum_is_idempotent(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)

    first = _upload(client, unit3.id)
    assert first.status_code == 201
    document_id = first.json()["id"]
    chunk_count = first.json()["chunk_count"]

    second = _upload(client, unit3.id)
    assert second.status_code == 201
    assert second.json()["id"] == document_id
    assert second.json()["chunk_count"] == chunk_count

    docs = seeded_db.scalars(select(KnowledgeDocument).where(KnowledgeDocument.unit_id == unit3.id)).all()
    assert len(docs) == 1


def test_upload_replaces_chunks_when_content_changes(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    first = _upload(client, unit3.id)
    assert first.status_code == 201
    document_id = first.json()["id"]

    modified_path = tmp_path / "modified.docx"
    doc = Document(str(KB_FILE))
    doc.paragraphs[0].add_run(" (updated)")
    doc.save(str(modified_path))

    resp = _upload(client, unit3.id, content=modified_path.read_bytes())

    assert resp.status_code == 201
    assert resp.json()["id"] == document_id
    assert resp.json()["chunk_count"] > 0


def test_upload_rejects_missing_unit(client, seeded_db):
    _login(client, seeded_db)

    resp = _upload(client, "00000000-0000-0000-0000-000000000000")

    assert resp.status_code == 400


def test_upload_rejects_non_docx_file(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)

    resp = client.post(
        "/admin/knowledge-documents",
        data={"unit_id": str(unit3.id)},
        files={"file": ("notes.txt", b"plain text", "text/plain")},
    )

    assert resp.status_code == 400


def test_list_documents_returns_all(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    unit2 = _unit2_grade7(seeded_db)
    _upload(client, unit3.id)
    _upload(client, unit2.id, filename="GS7 - UNIT 2 - LESSON.docx")

    resp = client.get("/admin/knowledge-documents")

    assert resp.status_code == 200
    assert len(resp.json()) == 2


def test_toggle_publish_status(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    document_id = _upload(client, unit3.id).json()["id"]

    resp = client.patch(f"/admin/knowledge-documents/{document_id}", json={"is_published": False})

    assert resp.status_code == 200
    assert resp.json()["is_published"] is False


def test_delete_document_cascades_chunks(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    document_id = _upload(client, unit3.id).json()["id"]

    resp = client.delete(f"/admin/knowledge-documents/{document_id}")

    assert resp.status_code == 204
    assert seeded_db.get(KnowledgeDocument, document_id) is None
    assert client.get("/admin/knowledge-documents").json() == []


def test_update_and_delete_not_found(client, seeded_db):
    _login(client, seeded_db)
    missing_id = "00000000-0000-0000-0000-000000000000"

    assert client.patch(f"/admin/knowledge-documents/{missing_id}", json={"is_published": False}).status_code == 404
    assert client.delete(f"/admin/knowledge-documents/{missing_id}").status_code == 404


def test_list_document_chunks(client, seeded_db):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    document_id = _upload(client, unit3.id).json()["id"]

    resp = client.get(f"/admin/knowledge-documents/{document_id}/chunks")

    assert resp.status_code == 200
    chunks = resp.json()
    assert len(chunks) > 0
    assert chunks == sorted(chunks, key=lambda c: c["order_no"])
    first = chunks[0]
    assert set(first.keys()) == {"id", "order_no", "chunk_type", "section_title", "raw_text", "structured"}


def test_list_document_chunks_not_found(client, seeded_db):
    _login(client, seeded_db)

    resp = client.get("/admin/knowledge-documents/00000000-0000-0000-0000-000000000000/chunks")

    assert resp.status_code == 404


def test_upload_grammar_reference_document(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    point = _present_simple_point(seeded_db)
    content = _build_grammar_reference_docx(tmp_path)

    resp = _upload_grammar(client, point.id, content=content)

    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["unit"] is None
    assert body["grammar_point"]["name"] == "Present Simple"
    assert body["grammar_point"]["topic_name"].startswith("Tense")
    assert body["chunk_count"] > 0


def test_upload_grammar_reference_keeps_table_chunk(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    point = _present_simple_point(seeded_db)
    content = _build_grammar_reference_docx(tmp_path)
    document_id = _upload_grammar(client, point.id, content=content).json()["id"]

    chunks = client.get(f"/admin/knowledge-documents/{document_id}/chunks").json()

    assert any("S + V…" in c["raw_text"] for c in chunks)
    assert all(c["chunk_type"] == "grammar" for c in chunks)


def test_upload_grammar_reference_is_idempotent(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    point = _present_simple_point(seeded_db)
    content = _build_grammar_reference_docx(tmp_path)

    first = _upload_grammar(client, point.id, content=content)
    second = _upload_grammar(client, point.id, content=content)

    assert first.json()["id"] == second.json()["id"]
    docs = seeded_db.scalars(
        select(KnowledgeDocument).where(KnowledgeDocument.grammar_point_id == point.id)
    ).all()
    assert len(docs) == 1


def test_upload_rejects_both_unit_and_grammar_point(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    point = _present_simple_point(seeded_db)
    content = _build_grammar_reference_docx(tmp_path)

    resp = client.post(
        "/admin/knowledge-documents",
        data={"unit_id": str(unit3.id), "grammar_point_id": str(point.id)},
        files={"file": ("x.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert resp.status_code == 400


def test_upload_rejects_neither_unit_nor_grammar_point(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    content = _build_grammar_reference_docx(tmp_path)

    resp = client.post(
        "/admin/knowledge-documents",
        data={},
        files={"file": ("x.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert resp.status_code == 400


def test_upload_rejects_missing_grammar_point(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    content = _build_grammar_reference_docx(tmp_path)

    resp = _upload_grammar(client, "00000000-0000-0000-0000-000000000000", content=content)

    assert resp.status_code == 400


def test_list_documents_returns_both_unit_and_grammar_sources(client, seeded_db, tmp_path):
    _login(client, seeded_db)
    unit3 = _unit3_grade7(seeded_db)
    point = _present_simple_point(seeded_db)
    _upload(client, unit3.id)
    _upload_grammar(client, point.id, content=_build_grammar_reference_docx(tmp_path))

    docs = client.get("/admin/knowledge-documents").json()

    assert len(docs) == 2
    assert {d["unit"] is None for d in docs} == {True, False}
