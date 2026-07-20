from docx import Document

from app.models.knowledge import DocumentChunkType
from app.services.grammar_parser import parse_grammar_reference_docx


def _bold_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True


def _build_reference_doc(tmp_path) -> str:
    doc = Document()
    _bold_paragraph(doc, "1. The Simple present (Thì hiện tại đơn)")
    _bold_paragraph(doc, "Cách dùng:")
    doc.add_paragraph("The earth is round. It goes around the sun.")
    doc.add_paragraph("She always gets up early.")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "be (+)"
    table.rows[0].cells[1].text = "S + am/is/are…"
    table.rows[1].cells[0].text = "V (+)"
    table.rows[1].cells[1].text = "S + V…"

    _bold_paragraph(doc, "Cách phát âm -s/-es")
    doc.add_paragraph("laugh /lɑːf/ -> laughs /lɑːfs/")

    path = tmp_path / "simple-present.docx"
    doc.save(str(path))
    return str(path)


def test_parses_bold_headers_and_groups_content_under_them(tmp_path):
    path = _build_reference_doc(tmp_path)

    chunks = parse_grammar_reference_docx(path)

    assert all(c.chunk_type == DocumentChunkType.GRAMMAR for c in chunks)

    usage_chunks = [c for c in chunks if c.section_title == "Cách dùng:"]
    assert any("earth is round" in c.raw_text for c in usage_chunks)
    assert any("gets up early" in c.raw_text for c in usage_chunks)

    pronunciation_chunks = [c for c in chunks if c.section_title == "Cách phát âm -s/-es"]
    assert any("laughs" in c.raw_text for c in pronunciation_chunks)


def test_keeps_tables_as_chunks_under_current_section(tmp_path):
    path = _build_reference_doc(tmp_path)

    chunks = parse_grammar_reference_docx(path)

    table_chunks = [c for c in chunks if "be (+)" in c.raw_text]
    assert len(table_chunks) == 1
    assert table_chunks[0].section_title == "Cách dùng:"
    assert "S + am/is/are…" in table_chunks[0].raw_text


def test_order_no_is_sequential_and_unique(tmp_path):
    path = _build_reference_doc(tmp_path)

    chunks = parse_grammar_reference_docx(path)

    order_nos = [c.order_no for c in chunks]
    assert order_nos == sorted(order_nos)
    assert len(order_nos) == len(set(order_nos))


def test_never_raises_when_document_has_no_bold_headers(tmp_path):
    doc = Document()
    doc.add_paragraph("Plain text with no formatting at all.")
    doc.add_paragraph("Another plain paragraph.")
    path = tmp_path / "no-headers.docx"
    doc.save(str(path))

    chunks = parse_grammar_reference_docx(str(path))

    assert len(chunks) == 2
    assert all(c.section_title == "" for c in chunks)
